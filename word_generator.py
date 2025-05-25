#!/usr/bin/env python
"""
Word document generator for CV data.
Creates a Microsoft Word document that can be easily edited.
"""
import os
import logging
from typing import Dict, Any, Optional, List

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT

from cv_data import CVData
from theme import Theme

# Configure logging
logger = logging.getLogger('word_generator')

class WordCVGenerator:
    """Generator for CV in Microsoft Word format."""
    
    def __init__(self, debug_mode: bool = False):
        """Initialize the Word CV generator."""
        self.debug_mode = debug_mode
        if debug_mode:
            logger.setLevel(logging.DEBUG)
    
    def create_cv(self, data_path: str, output_path: str) -> None:
        """
        Create a Word document CV.
        
        Args:
            data_path: Path to JSON data file
            output_path: Path to save the Word document
        """
        # Ensure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Load CV data
        cv_data = CVData(self.debug_mode)
        cv_data.load(data_path, debug_mode=self.debug_mode)
        cv_data.validate()
        
        # Create the document
        doc = Document()
        
        # Set up the document
        self._setup_document(doc, cv_data)
        
        # Add content sections
        self._add_header(doc, cv_data)
        
        # Left and right columns
        self._add_two_column_layout(doc, cv_data)
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Word document created successfully: {output_path}")
    
    def _setup_document(self, doc: Document, cv_data: CVData) -> None:
        """Configure document settings."""
        # Get layout data from CVData
        layout_data = cv_data.get_layout_data()

        # Set page size based on layout settings
        page_size = layout_data.get("page_size", "A4").upper()
        
        # Set margins using values from layout_data (converting inches to Inches for docx)
        sections = doc.sections
        for section in sections:
            # Set page size
            if page_size == "A4":
                section.page_width = Inches(8.27)  # A4 width in inches
                section.page_height = Inches(11.69)  # A4 height in inches
            elif page_size == "LETTER":
                section.page_width = Inches(8.5)  # Letter width in inches
                section.page_height = Inches(11)  # Letter height in inches
            elif page_size == "LEGAL":
                section.page_width = Inches(8.5)  # Legal width in inches
                section.page_height = Inches(14)  # Legal height in inches
            elif page_size == "A3":
                section.page_width = Inches(11.69)  # A3 width in inches
                section.page_height = Inches(16.54)  # A3 height in inches
                
            # Set margins
            section.top_margin = Inches(layout_data.get("top_margin", 0.4))  # Use consistent defaults
            section.bottom_margin = Inches(layout_data.get("bottom_margin", 0.4))
            section.left_margin = Inches(layout_data.get("left_margin", 0.3))
            section.right_margin = Inches(layout_data.get("right_margin", 0.3))
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def _add_header(self, doc: Document, cv_data: CVData) -> None:
        """Add the header with name and contact info."""
        candidate_data = cv_data.get_candidate_data()
        
        # Add name as heading
        name = candidate_data.get("name", "")
        heading = doc.add_heading(name, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_format = heading.runs[0].font
        heading_format.size = Pt(18)
        heading_format.bold = True
        heading_format.color.rgb = RGBColor(0, 48, 135)  # Dark blue
        
        # Add contact info in a table (3 columns)
        contact_details = candidate_data.get("contact", [])
        if contact_details:
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.style = 'Table Grid'
            
            # Make table borders invisible - safer approach without using tcBorders
            for cell in table.rows[0].cells:
                cell._element.get_or_add_tcPr()
                cell.border_bottom = cell.border_top = cell.border_left = cell.border_right = None
            
            # Distribute contact info across columns
            contacts_per_col = (len(contact_details) + 2) // 3
            row_cells = table.rows[0].cells
            
            for i, col in enumerate(range(3)):
                start_idx = col * contacts_per_col
                end_idx = min(start_idx + contacts_per_col, len(contact_details))
                for contact in contact_details[start_idx:end_idx]:
                    text = contact.get("text", "")
                    p = row_cells[i].add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(text)
                    run.font.size = Pt(9)
                    
        # Add a line to separate header from content
        doc.add_paragraph("_" * 150)
    
    def _add_two_column_layout(self, doc: Document, cv_data: CVData) -> None:
        """
        Add content in a two-column layout. 
        In Word, we'll use a table with two columns to simulate the two-column layout.
        """
        table = doc.add_table(rows=1, cols=2)
        
        # Get layout data for consistent column widths
        layout_data = cv_data.get_layout_data()
        left_column_ratio = layout_data.get("left_column_width_ratio", 0.3)
        
        # Determine page width based on page size
        page_size = layout_data.get("page_size", "A4").upper()
        if page_size == "A4":
            page_width = 8.27  # A4 width in inches
        elif page_size == "LETTER":
            page_width = 8.5  # Letter width in inches
        elif page_size == "LEGAL":
            page_width = 8.5  # Legal width in inches
        elif page_size == "A3":
            page_width = 11.69  # A3 width in inches
        else:
            page_width = 8.27  # Default to A4
        
        # Calculate available width accounting for margins
        available_width = page_width - layout_data.get("left_margin", 0.3) - layout_data.get("right_margin", 0.3)
        
        # Set column widths (using left_column_width_ratio from layout)
        left_column_width = available_width * left_column_ratio
        right_column_width = available_width - left_column_width
        
        table.columns[0].width = Inches(left_column_width)
        table.columns[1].width = Inches(right_column_width)
        
        # Remove table borders - safer approach
        for cell in table.rows[0].cells:
            cell._element.get_or_add_tcPr()
            cell.border_bottom = cell.border_top = cell.border_left = cell.border_right = None
        
        # Get cells for left and right columns
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        # Set light background for left column using XML approach
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # Create shading element with light blue color
        shading_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F8FC"/>')
        left_cell._element.get_or_add_tcPr().append(shading_xml)
        
        # Add content to left column
        self._add_left_column_content(left_cell, cv_data)
        
        # Add content to right column
        self._add_right_column_content(right_cell, cv_data)
    
    def _add_left_column_content(self, cell: Any, cv_data: CVData) -> None:
        """Add content to the left column."""
        # Technical Skills section
        if cv_data.has_section("technical_skills"):
            skills_data = cv_data.get_technical_skills()
            
            # Add section header
            p = cell.add_paragraph()
            run = p.add_run("TECHNICAL EXPERTISE")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add skills by category
            for category, skills in skills_data.items():
                p = cell.add_paragraph()
                run = p.add_run(category)
                run.bold = True
                run.font.size = Pt(10)
                
                for skill in skills:
                    p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Pt(10)
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run(f"• {skill}")
                    run.font.size = Pt(9.5)
        
        # Education section
        if cv_data.has_section("education"):
            education_data = cv_data.get_education_data()
            
            # Add section header with some space before
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("EDUCATION")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add education items
            for edu in education_data.get("items", []):
                institution = edu.get("institution", "")
                degree = edu.get("degree", "")
                duration = edu.get("duration", "")
                
                # Institution name
                p = cell.add_paragraph()
                run = p.add_run(institution)
                run.bold = True
                run.font.size = Pt(10)
                
                # Degree
                if degree:
                    p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Pt(10)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(degree)
                    run.font.size = Pt(9.5)
                
                # Duration
                if duration:
                    p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Pt(10)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(duration)
                    run.font.size = Pt(9.5)
        
        # More details section
        if cv_data.has_section("additional_info"):
            additional_info = cv_data.get_additional_info()
            
            # Add section header with some space before
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("MORE DETAILS")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add additional info
            for info in additional_info:
                p = cell.add_paragraph()
                p.paragraph_format.left_indent = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(f"• {info}")
                run.font.size = Pt(9.5)
    
    def _add_right_column_content(self, cell: Any, cv_data: CVData) -> None:
        """Add content to the right column."""
        # Profile section
        profile_text = cv_data.get_profile_data()
        if profile_text:
            # Add section header
            p = cell.add_paragraph()
            run = p.add_run("PROFILE")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add profile text
            paragraphs = profile_text.split('\n\n')
            for paragraph in paragraphs:
                p = cell.add_paragraph()
                p.paragraph_format.left_indent = Pt(5)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(paragraph)
                run.font.size = Pt(9.5)
                # Set paragraph justification
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Professional Experience section
        experience_data = cv_data.get_experience_data()
        if experience_data and experience_data.get("companies"):
            # Add section header with space before
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("PROFESSIONAL EXPERIENCE")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add companies
            companies = experience_data.get("companies", [])
            for company in companies:
                company_name = company.get("name", "")
                duration = company.get("totalDuration", "")
                
                # Company name and duration
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(f"{company_name} | {duration}")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 48, 135)  # Dark blue
                
                # Roles
                for role in company.get("roles", []):
                    role_title = role.get("title", "")
                    responsibilities = role.get("responsibilities", [])
                    
                    # Role title
                    p = cell.add_paragraph()
                    p.paragraph_format.left_indent = Pt(10)
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run(role_title)
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0, 112, 204)  # Lighter blue
                    
                    # Responsibilities
                    for resp in responsibilities:
                        p = cell.add_paragraph()
                        p.paragraph_format.left_indent = Pt(15)
                        p.paragraph_format.space_after = Pt(3)
                        p.style = 'List Bullet'
                        run = p.add_run(resp)
                        run.font.size = Pt(9)
        
        # Projects section
        projects_data = cv_data.get_projects_data()
        if projects_data:
            # Add section header with space before
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("PROJECTS & ACHIEVEMENTS")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add projects
            for project in projects_data:
                title = project.get("title", "")
                description = project.get("description", "")
                
                # Project title
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 48, 135)  # Dark blue
                
                # Project description
                p = cell.add_paragraph()
                p.paragraph_format.left_indent = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(description)
                run.font.size = Pt(9)
                # Set paragraph justification
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # References
        references = cv_data.get_references()
        if references:
            # Add section header with space before
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run("REFERENCES")
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            # Add references text
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Pt(5)
            run = p.add_run(references)
            run.font.size = Pt(9)

# Command-line execution
if __name__ == "__main__":
    import argparse
    import sys
    
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[logging.StreamHandler(sys.stdout)]
    )
    
    # Parse arguments
    parser = argparse.ArgumentParser(description='Generate CV in Microsoft Word format')
    parser.add_argument('data_path', help='Path to CV data JSON file')
    parser.add_argument('output_path', help='Path to save the Word document')
    parser.add_argument('--debug', action='store_true', help='Enable debug mode')
    
    args = parser.parse_args()
    
    # Generate Word document
    generator = WordCVGenerator(debug_mode=args.debug)
    try:
        generator.create_cv(args.data_path, args.output_path)
    except Exception as e:
        logger.error(f"Error generating Word document: {e}")
        sys.exit(1) 